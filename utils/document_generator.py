"""
Générateur de documents (DOCX, PDF) pour l'export des candidatures
"""

from datetime import datetime
from pathlib import Path
import io


def generate_cv_docx(cv_content: str, nom: str = "Valerie_Jasica") -> io.BytesIO:
    """
    Génère un fichier DOCX à partir du contenu du CV.
    
    Args:
        cv_content: Contenu du CV en texte/markdown
        nom: Nom pour le fichier
        
    Returns:
        BytesIO contenant le fichier DOCX
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Titre
        title = doc.add_heading('CV - ' + nom.replace('_', ' '), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date de génération
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()
        
        # Contenu
        for line in cv_content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        # Sauvegarder dans un buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except ImportError:
        raise ImportError("python-docx n'est pas installé. Exécutez: pip install python-docx")


def generate_lettre_docx(lettre_content: str, nom: str = "Valerie_Jasica") -> io.BytesIO:
    """
    Génère un fichier DOCX à partir de la lettre de motivation.
    
    Args:
        lettre_content: Contenu de la lettre
        nom: Nom pour le fichier
        
    Returns:
        BytesIO contenant le fichier DOCX
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # Style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # En-tête avec coordonnées
        header_para = doc.add_paragraph()
        header_para.add_run("Valérie JASICA\n").bold = True
        header_para.add_run("COULLONS 45720\n")
        header_para.add_run("06 09 91 41 86\n")
        header_para.add_run("valerie.jasica67@gmail.com")
        
        # Date à droite
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f"Le {datetime.now().strftime('%d %B %Y')}")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Contenu de la lettre
        for line in lettre_content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            else:
                # Nettoyer le markdown basique
                clean_line = line.replace('**', '')
                doc.add_paragraph(clean_line)
        
        # Signature
        doc.add_paragraph()
        doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_para.add_run("Valérie JASICA")
        
        # Sauvegarder dans un buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except ImportError:
        raise ImportError("python-docx n'est pas installé. Exécutez: pip install python-docx")


def generate_preparation_docx(preparation_content: str, poste: str = "Poste") -> io.BytesIO:
    """
    Génère un fichier DOCX à partir de la préparation d'entretien.
    
    Args:
        preparation_content: Contenu de la préparation
        poste: Titre du poste pour le document
        
    Returns:
        BytesIO contenant le fichier DOCX
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Marges
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Titre
        title = doc.add_heading(f'Préparation Entretien - {poste}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f"Préparé le {datetime.now().strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()
        
        # Contenu
        for line in preparation_content.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            elif line.startswith('- ') or line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('> '):
                # Citation / question
                p = doc.add_paragraph()
                p.add_run(line[2:]).italic = True
            else:
                doc.add_paragraph(line.replace('**', ''))
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except ImportError:
        raise ImportError("python-docx n'est pas installé")


def generate_pdf_from_html(content: str, title: str = "Document") -> io.BytesIO:
    """
    Génère un PDF à partir de contenu HTML/Markdown.
    Utilise une approche simple avec reportlab.
    
    Args:
        content: Contenu texte/markdown
        title: Titre du document
        
    Returns:
        BytesIO contenant le PDF
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        
        buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        # Styles personnalisés
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=20
        ))
        
        styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=15
        ))
        
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=6
        ))
        
        styles.add(ParagraphStyle(
            name='CustomBullet',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=4
        ))
        
        # Construire le contenu
        story = []
        
        # Titre
        story.append(Paragraph(title, styles['CustomTitle']))
        story.append(Spacer(1, 12))
        
        # Parser le contenu
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            elif line.startswith('###'):
                text = line.replace('###', '').strip()
                story.append(Paragraph(text, styles['Heading3']))
            elif line.startswith('##'):
                text = line.replace('##', '').strip()
                story.append(Paragraph(text, styles['CustomHeading']))
            elif line.startswith('#'):
                text = line.replace('#', '').strip()
                story.append(Paragraph(text, styles['Heading1']))
            elif line.startswith('- ') or line.startswith('• '):
                text = '• ' + line[2:].replace('**', '<b>', 1).replace('**', '</b>', 1)
                story.append(Paragraph(text, styles['CustomBullet']))
            else:
                # Nettoyer le markdown
                text = line.replace('**', '<b>', 1).replace('**', '</b>', 1)
                text = text.replace('*', '<i>', 1).replace('*', '</i>', 1)
                try:
                    story.append(Paragraph(text, styles['CustomBody']))
                except:
                    # Si erreur de parsing, utiliser texte brut
                    story.append(Paragraph(line.replace('**', '').replace('*', ''), styles['CustomBody']))
        
        doc.build(story)
        buffer.seek(0)
        
        return buffer
        
    except ImportError:
        raise ImportError("reportlab n'est pas installé")


def generate_all_documents(
    cv_content: str = None,
    lettre_content: str = None,
    preparation_content: str = None,
    poste: str = "Poste",
    entreprise: str = "Entreprise"
) -> dict:
    """
    Génère tous les documents disponibles en DOCX.
    
    Args:
        cv_content: Contenu du CV adapté
        lettre_content: Contenu de la lettre de motivation
        preparation_content: Contenu de la préparation d'entretien
        poste: Titre du poste
        entreprise: Nom de l'entreprise
        
    Returns:
        Dict avec les buffers des fichiers générés
    """
    documents = {}
    
    if cv_content:
        try:
            documents['cv_docx'] = generate_cv_docx(cv_content)
        except Exception as e:
            print(f"Erreur génération CV DOCX: {e}")
    
    if lettre_content:
        try:
            documents['lettre_docx'] = generate_lettre_docx(lettre_content)
        except Exception as e:
            print(f"Erreur génération lettre DOCX: {e}")
    
    if preparation_content:
        try:
            documents['preparation_docx'] = generate_preparation_docx(preparation_content, poste)
        except Exception as e:
            print(f"Erreur génération préparation DOCX: {e}")
    
    return documents


def save_to_file(content: str, filename: str, output_dir: str = "exports") -> Path:
    """
    Sauvegarde du contenu dans un fichier texte.
    
    Args:
        content: Contenu à sauvegarder
        filename: Nom du fichier
        output_dir: Répertoire de sortie
        
    Returns:
        Path vers le fichier créé
    """
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    file_path = output_path / filename
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return file_path

